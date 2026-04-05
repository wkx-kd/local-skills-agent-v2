"""
文件解析服务 — 多格式文件文本提取

支持格式：
- PDF: pdfplumber (文本提取) + pypdf (备选)
- Word: python-docx
- Excel: openpyxl / pandas
- CSV: pandas
- 图片: 暂不支持 OCR，返回提示
- 代码/文本: 直接读取
"""

import os
from pathlib import Path
from typing import Optional
from dataclasses import dataclass


@dataclass
class ParsedFile:
    """解析后的文件内容"""
    filename: str
    file_type: str
    text: str
    char_count: int
    page_count: Optional[int] = None  # PDF 页数
    error: Optional[str] = None


class FileParser:
    """多格式文件解析器"""

    # 支持的文件类型
    TEXT_EXTENSIONS = {'.txt', '.md', '.json', '.yaml', '.yml', '.xml', '.html', '.css'}
    CODE_EXTENSIONS = {'.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.c', '.cpp', '.go', '.rs', '.rb', '.php', '.sh', '.sql'}
    PDF_EXTENSIONS = {'.pdf'}
    WORD_EXTENSIONS = {'.docx', '.doc'}
    EXCEL_EXTENSIONS = {'.xlsx', '.xls'}
    CSV_EXTENSIONS = {'.csv'}
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}

    # 小文件阈值（字符数），小于此值标记为 full_text 策略
    SMALL_FILE_THRESHOLD = 20000  # 约 5000 tokens

    @classmethod
    def parse(cls, file_path: Path) -> ParsedFile:
        """
        解析文件，提取文本内容。

        Args:
            file_path: 文件绝对路径

        Returns:
            ParsedFile 对象
        """
        if not file_path.exists():
            return ParsedFile(
                filename=file_path.name,
                file_type='unknown',
                text='',
                char_count=0,
                error='文件不存在',
            )

        suffix = file_path.suffix.lower()
        filename = file_path.name

        try:
            # 根据扩展名选择解析方法
            if suffix in cls.PDF_EXTENSIONS:
                return cls._parse_pdf(file_path)
            elif suffix in cls.WORD_EXTENSIONS:
                return cls._parse_word(file_path)
            elif suffix in cls.EXCEL_EXTENSIONS:
                return cls._parse_excel(file_path)
            elif suffix in cls.CSV_EXTENSIONS:
                return cls._parse_csv(file_path)
            elif suffix in cls.IMAGE_EXTENSIONS:
                return ParsedFile(
                    filename=filename,
                    file_type='image',
                    text='[图片文件，暂不支持文本提取]',
                    char_count=0,
                )
            elif suffix in cls.TEXT_EXTENSIONS or suffix in cls.CODE_EXTENSIONS:
                return cls._parse_text(file_path, is_code=suffix in cls.CODE_EXTENSIONS)
            else:
                # 尝试作为文本读取
                return cls._parse_text(file_path)

        except Exception as e:
            return ParsedFile(
                filename=filename,
                file_type=suffix.lstrip('.') or 'unknown',
                text='',
                char_count=0,
                error=str(e),
            )

    @classmethod
    def _parse_text(cls, file_path: Path, is_code: bool = False) -> ParsedFile:
        """解析文本/代码文件"""
        try:
            content = file_path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            content = file_path.read_text(encoding='gbk')

        file_type = 'code' if is_code else file_path.suffix.lstrip('.').lower()
        return ParsedFile(
            filename=file_path.name,
            file_type=file_type or 'text',
            text=content,
            char_count=len(content),
        )

    @classmethod
    def _parse_pdf(cls, file_path: Path) -> ParsedFile:
        """解析 PDF 文件"""
        import pdfplumber

        text_parts = []
        page_count = 0

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n\n".join(text_parts)
        return ParsedFile(
            filename=file_path.name,
            file_type='pdf',
            text=full_text,
            char_count=len(full_text),
            page_count=page_count,
        )

    @classmethod
    def _parse_word(cls, file_path: Path) -> ParsedFile:
        """解析 Word 文件"""
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)
        return ParsedFile(
            filename=file_path.name,
            file_type='docx',
            text=full_text,
            char_count=len(full_text),
        )

    @classmethod
    def _parse_excel(cls, file_path: Path) -> ParsedFile:
        """解析 Excel 文件"""
        import pandas as pd

        # 读取所有 sheet
        sheets = pd.read_excel(file_path, sheet_name=None)
        text_parts = []

        for sheet_name, df in sheets.items():
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            # 将 DataFrame 转为文本
            text_parts.append(df.to_string(index=False, na_rep=''))
            text_parts.append('')

        full_text = "\n".join(text_parts)
        return ParsedFile(
            filename=file_path.name,
            file_type='xlsx',
            text=full_text,
            char_count=len(full_text),
        )

    @classmethod
    def _parse_csv(cls, file_path: Path) -> ParsedFile:
        """解析 CSV 文件"""
        import pandas as pd

        df = pd.read_csv(file_path)
        full_text = df.to_string(index=False, na_rep='')

        return ParsedFile(
            filename=file_path.name,
            file_type='csv',
            text=full_text,
            char_count=len(full_text),
        )

    @classmethod
    def get_processing_strategy(cls, char_count: int) -> str:
        """
        根据文本长度决定处理策略。

        Args:
            char_count: 字符数

        Returns:
            'full_text' 或 'rag'
        """
        return 'full_text' if char_count < cls.SMALL_FILE_THRESHOLD else 'rag'

    @classmethod
    def chunk_text(cls, text: str, chunk_size: int = 512, overlap: int = 50) -> list[dict]:
        """
        将长文本分块。

        Args:
            text: 原始文本
            chunk_size: 每块最大字符数
            overlap: 块之间的重叠字符数

        Returns:
            [{"index": 0, "content": "..."}, ...]
        """
        if len(text) <= chunk_size:
            return [{"index": 0, "content": text}]

        chunks = []
        start = 0
        index = 0

        while start < len(text):
            end = start + chunk_size
            chunk_content = text[start:end]

            # 尝试在句子边界截断
            if end < len(text):
                # 查找最后一个句号/问号/感叹号
                last_sentence_end = max(
                    chunk_content.rfind('。'),
                    chunk_content.rfind('？'),
                    chunk_content.rfind('！'),
                    chunk_content.rfind('.'),
                    chunk_content.rfind('?'),
                    chunk_content.rfind('!'),
                )
                if last_sentence_end > chunk_size // 2:
                    chunk_content = chunk_content[:last_sentence_end + 1]
                    end = start + last_sentence_end + 1

            chunks.append({"index": index, "content": chunk_content.strip()})
            index += 1
            start = end - overlap if end < len(text) else end

        return chunks